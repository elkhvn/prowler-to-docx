import json
import argparse
import collections
import docx
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT


def parse_args():
    parser = argparse.ArgumentParser(description='Process some JSON data.')
    parser.add_argument('filename', type=str, help='The filename of the JSON file')
    return parser.parse_args()



def load_json(filename):
    """ Load and return the JSON data from a file """
    try:
        with open(filename, 'r') as file:
            return json.load(file)
    except FileNotFoundError:
        print(f"Error: The file '{filename}' does not exist.")
        exit(1)
    except json.JSONDecodeError:
        print("Error: Failed to decode JSON from the file.")
        exit(1)


def nested_dict(n, type):
    if n == 1:
        return collections.defaultdict(type)
    else:
        return collections.defaultdict(lambda: nested_dict(n-1, type))



def write_grouped_data_to_doc(grouped, doc):

    styles = doc.styles

    # Create a custom heading style
    heading_style = styles['Heading 2']
    heading_style.font.name = 'Rockwell'
    heading_style.font.size = Pt(18)
    heading_style.font.color.rgb = RGBColor(187, 165, 61)  # Chrome Gold



    subheading_style = styles['Heading 3']
    subheading_style.font.name = 'Rockwell'
    subheading_style.font.size = Pt(15)
    subheading_style.font.color.rgb = RGBColor(0, 102, 204)  # Blue


    for service_name, check_ids in grouped.items():
        service_name = ' '.join(service_name.split('-')).upper()
        # Add a paragraph with the custom style
        doc.add_paragraph(service_name, style='Heading 2')




        for check_id, findings in check_ids.items():
            check_id = ' '.join(check_id.split('_')).upper()
            # Add a paragraph with the custom style
            doc.add_paragraph(check_id, style='Heading 3')


            status_group = collections.defaultdict(list)    
            for finding in findings:
                status_ext = finding['status_ext']
                arn = finding['arn']

                status_group[status_ext].append(arn)



            # ================== DESCRIPTION ======================

            # Create a new paragraph and justify it
            paragraph = doc.add_paragraph()
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

            # Add a bold run to the paragraph 
            run = paragraph.add_run('Description: ')
            run.bold = True
            run.font.name = 'Rockwell'  # Setting font (if you want the same font for all runs)
            run.font.size = Pt(12)      # Setting font size (if desired)

            # Add another run for the URL
            run = paragraph.add_run(finding['description'])
            run.font.name = 'Rockwell'
            run.font.size = Pt(12)  # Assuming you want the same size, adjust as necessary





            # ================= REMEDIATION =========================
            # Create a new paragraph and justify it
            paragraph = doc.add_paragraph()
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

            # Add a bold run to the paragraph 
            run = paragraph.add_run('Remediation: ')
            run.bold = True
            run.font.name = 'Rockwell'  # Setting font (if you want the same font for all runs)
            run.font.size = Pt(12)      # Setting font size (if desired)

            # Add another run for the URL
            run = paragraph.add_run(finding['remediation_text'])
            run.font.name = 'Rockwell'
            run.font.size = Pt(12)  # Assuming you want the same size, adjust as necessary




            # ================ CODE =================================
            # Create a new paragraph and justify it
            paragraph = doc.add_paragraph()
            # Add another run for the URL
            run = paragraph.add_run(finding['remediation_code'])
            run.font.name = 'Courier New'
            run.font.size = Pt(12)  # Assuming you want the same size, adjust as necessary




            # ================ REFERENCE =================================
            # Create a new paragraph and justify it
            paragraph = doc.add_paragraph()
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

            # Add a bold run to the paragraph for "Reference:"
            run = paragraph.add_run('Reference:\n')
            run.bold = True
            run.font.name = 'Rockwell'  # Setting font (if you want the same font for all runs)
            run.font.size = Pt(12)      # Setting font size (if desired)

            # Add another run for the URL
            run = paragraph.add_run(finding['remediation_url'])
            run.font.name = 'Rockwell'
            run.font.size = Pt(12)  # Assuming you want the same size, adjust as necessary




            # ================ RISK =================================
            # Create a new paragraph and justify it
            paragraph = doc.add_paragraph()
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

            # Add a bold run to the paragraph 
            run = paragraph.add_run('Risk: ')
            run.bold = True
            run.font.name = 'Rockwell'  # Setting font (if you want the same font for all runs)
            run.font.size = Pt(12)      # Setting font size (if desired)

            # Add another run for the URL
            run = paragraph.add_run(finding['risk'])
            run.font.name = 'Rockwell'
            run.font.size = Pt(12)  # Assuming you want the same size, adjust as necessary




            # ================ SEVERITY =================================
            # Create a new paragraph and justify it
            paragraph = doc.add_paragraph()
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

            # Add a bold run to the paragraph 
            run = paragraph.add_run('Severity: ')
            run.bold = True
            run.font.name = 'Rockwell'  # Setting font (if you want the same font for all runs)
            run.font.size = Pt(12)      # Setting font size (if desired)

            # Add another run for the URL
            run = paragraph.add_run(finding['severity'])
            run.font.name = 'Rockwell'
            run.font.size = Pt(12)  # Assuming you want the same size, adjust as necessary



            
            for status_extended, arns in status_group.items():
                paragraph = doc.add_paragraph()
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                run = paragraph.add_run(status_extended)
                run.bold = True
                run.font.name = 'Rockwell'
                run.font.size = Pt(12)      # Setting font size (if desired)

                for arn in arns:
                    paragraph = doc.add_paragraph(style='List Bullet')
                    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                    run = paragraph.add_run(arn)
                    run.font.name = 'Rockwell'
        


            # Add a paragraph
            paragraph = doc.add_paragraph()
            paragraph.add_run('\n\n')
        
        
            print(f"[+] One check finished {check_id}")        
            
            
            
        # Add a paragraph
        paragraph = doc.add_paragraph()
        paragraph.add_run('\n\n\n')






def main():

    args = parse_args()



    json_data = load_json(args.filename) # load JSON data from the file



    grouped_data = nested_dict(2, list)  # 8 levels deep: ServiceName, CheckID, StatusExtended


    for item in json_data:
        service_name = item['ServiceName']
        status = item['Status']
        check_id = item['CheckID']



        findings = {
            'description': item['Description'],
            'remediation_text': item['Remediation']['Recommendation']['Text'],
            'remediation_code': item['Remediation']['Code']['CLI'],
            'remediation_url':  item['Remediation']['Recommendation']['Url'],
            'risk': item['Risk'],
            'severity': item['Severity'],
            'status_ext': item['StatusExtended'],
            'arn': item['ResourceArn']
        }


        
        if status == 'FAIL':
            if findings not in grouped_data[service_name][check_id]:
                grouped_data[service_name][check_id].append(findings)



    # create a Document 
    doc = docx.Document()
    write_grouped_data_to_doc(grouped_data, doc)


    # Save the document
    doc.save(args.filename.split('.')[0] + '.docx')


if __name__ == '__main__':
    main()

